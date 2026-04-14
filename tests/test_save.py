"""Tests for RevisionDocument.save accepting paths and file-like objects."""

import io
from pathlib import Path

import pytest
from docx import Document

from docx_revisions import RevisionDocument


class DescribeRevisionDocument_save_targets:
    """save() should accept str paths, Path objects, and binary streams."""

    def it_saves_to_bytesio_and_roundtrips(self):
        doc = Document()
        doc.add_paragraph("Hello BytesIO")
        rdoc = RevisionDocument(doc)

        buffer = io.BytesIO()
        rdoc.save(buffer)

        assert buffer.tell() > 0
        buffer.seek(0)

        rdoc2 = RevisionDocument(buffer)
        texts = [p.text for p in rdoc2.document.paragraphs]
        assert "Hello BytesIO" in texts

    def it_saves_to_string_path(self, tmp_path: Path):
        rdoc = RevisionDocument()
        path = tmp_path / "saved_str.docx"
        rdoc.save(str(path))
        assert path.exists()

        rdoc2 = RevisionDocument(str(path))
        assert rdoc2.document is not None

    def it_saves_to_path_object(self, tmp_path: Path):
        rdoc = RevisionDocument()
        path = tmp_path / "saved_path.docx"
        rdoc.save(path)
        assert path.exists()

        rdoc2 = RevisionDocument(str(path))
        assert rdoc2.document is not None

    def it_rejects_empty_string_path(self):
        rdoc = RevisionDocument()
        with pytest.raises(ValueError, match="must not be empty"):
            rdoc.save("")

    def it_rejects_non_path_non_stream(self):
        rdoc = RevisionDocument()
        with pytest.raises(TypeError):
            rdoc.save(123)  # type: ignore[arg-type]

        with pytest.raises(TypeError):
            rdoc.save(object())  # type: ignore[arg-type]

    def it_rejects_text_mode_stream(self, tmp_path: Path):
        rdoc = RevisionDocument()
        path = tmp_path / "text.docx"
        with open(path, "w") as text_stream, pytest.raises(ValueError, match="binary-mode stream"):
            rdoc.save(text_stream)  # type: ignore[arg-type]

    def it_saves_to_binary_file_handle(self, tmp_path: Path):
        rdoc = RevisionDocument()
        path = tmp_path / "binary.docx"
        with open(path, "wb") as binary_stream:
            rdoc.save(binary_stream)
        assert path.stat().st_size > 0
