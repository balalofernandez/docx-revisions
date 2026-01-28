"""Tests for docx_revisions core functionality."""

from datetime import datetime

import pytest
from docx import Document

from docx_revisions import (
    Revision,
    delete_with_tracking,
    get_accepted_text,
    get_revisions,
    insert_with_tracking,
    replace_with_tracking,
)


@pytest.fixture
def empty_paragraph():
    """Create a document with an empty paragraph."""
    doc = Document()
    return doc.add_paragraph("")


@pytest.fixture
def simple_paragraph():
    """Create a document with a simple text paragraph."""
    doc = Document()
    return doc.add_paragraph("Hello world")


class TestGetAcceptedText:
    def test_simple_text(self, simple_paragraph):
        result = get_accepted_text(simple_paragraph)
        assert result == "Hello world"

    def test_empty_paragraph(self, empty_paragraph):
        result = get_accepted_text(empty_paragraph)
        assert result == ""

    def test_multiple_runs(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hello ")
        p.add_run("world")
        result = get_accepted_text(p)
        assert result == "Hello world"


class TestGetRevisions:
    def test_no_revisions(self, simple_paragraph):
        revisions = get_revisions(simple_paragraph)
        assert revisions == []

    def test_empty_paragraph(self, empty_paragraph):
        revisions = get_revisions(empty_paragraph)
        assert revisions == []


class TestInsertWithTracking:
    def test_insert_at_end(self, simple_paragraph):
        insert_with_tracking(simple_paragraph, " today", author="Test")
        text = get_accepted_text(simple_paragraph)
        assert text == "Hello world today"

    def test_insert_creates_revision(self, simple_paragraph):
        insert_with_tracking(simple_paragraph, " today", author="TestAuthor")
        revisions = get_revisions(simple_paragraph)
        assert len(revisions) == 1
        assert revisions[0].type == "ins"
        assert revisions[0].text == " today"
        assert revisions[0].author == "TestAuthor"

    def test_insert_with_custom_date(self, empty_paragraph):
        custom_date = datetime(2024, 1, 15, 10, 30, 0)
        insert_with_tracking(empty_paragraph, "test", author="Author", date=custom_date)
        revisions = get_revisions(empty_paragraph)
        assert len(revisions) == 1
        assert revisions[0].date is not None


class TestDeleteWithTracking:
    def test_delete_word(self, simple_paragraph):
        delete_with_tracking(simple_paragraph, 6, 11, author="Test")  # Delete "world"
        text = get_accepted_text(simple_paragraph)
        assert text == "Hello "

    def test_delete_creates_revision(self, simple_paragraph):
        delete_with_tracking(simple_paragraph, 6, 11, author="TestAuthor")
        revisions = get_revisions(simple_paragraph)
        assert len(revisions) == 1
        assert revisions[0].type == "del"
        assert revisions[0].text == "world"
        assert revisions[0].author == "TestAuthor"


class TestReplaceWithTracking:
    def test_replace_word(self, simple_paragraph):
        result = replace_with_tracking(simple_paragraph, "world", "universe", author="Test")
        assert result is True
        text = get_accepted_text(simple_paragraph)
        assert text == "Hello universe"

    def test_replace_not_found(self, simple_paragraph):
        result = replace_with_tracking(simple_paragraph, "notfound", "replacement")
        assert result is False

    def test_replace_creates_two_revisions(self, simple_paragraph):
        replace_with_tracking(simple_paragraph, "world", "universe", author="Test")
        revisions = get_revisions(simple_paragraph)
        assert len(revisions) == 2
        types = {r.type for r in revisions}
        assert types == {"ins", "del"}


class TestRevisionModel:
    def test_revision_creation(self):
        rev = Revision(type="ins", text="hello", author="John")
        assert rev.type == "ins"
        assert rev.text == "hello"
        assert rev.author == "John"
        assert rev.date is None

    def test_revision_with_date(self):
        dt = datetime(2024, 1, 1)
        rev = Revision(type="del", text="removed", author="Jane", date=dt)
        assert rev.date == dt
