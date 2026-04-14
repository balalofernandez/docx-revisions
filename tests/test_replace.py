"""Tests for replace_tracked and replace_tracked_at."""

import pytest
from docx import Document

import docx_revisions  # noqa: F401
from docx_revisions import RevisionDocument, RevisionParagraph


class DescribeRevisionParagraph_replace_tracked:
    """Tests for RevisionParagraph.replace_tracked."""

    def it_replaces_within_single_run(self):
        doc = Document()
        para = doc.add_paragraph("Hello Unisys World")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("Unisys", "test", author="Tester")

        assert count == 1
        assert rp.has_track_changes is True
        # The deletions should contain "Unisys"
        assert any(d.text == "Unisys" for d in rp.deletions)
        # The insertions should contain "test"
        assert any(i.text == "test" for i in rp.insertions)

    def it_handles_multiple_occurrences(self):
        doc = Document()
        para = doc.add_paragraph("Unisys and Unisys again")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("Unisys", "test", author="Tester")

        assert count == 2
        assert len(rp.deletions) == 2
        assert len(rp.insertions) == 2

    def it_returns_zero_when_no_match(self):
        doc = Document()
        para = doc.add_paragraph("No match here")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("Unisys", "test", author="Tester")

        assert count == 0
        assert rp.has_track_changes is False

    def it_preserves_surrounding_text(self):
        doc = Document()
        para = doc.add_paragraph("Before Unisys After")
        rp = RevisionParagraph.from_paragraph(para)

        rp.replace_tracked("Unisys", "test", author="Tester")

        # accepted_text should have "Before test After"
        assert "Before" in rp.accepted_text
        assert "After" in rp.accepted_text
        assert "test" in rp.accepted_text

    def it_finds_text_spanning_multiple_runs(self):
        """Search text split across runs (e.g. OOXML run splitting)."""
        doc = Document()
        para = doc.add_paragraph("")
        # Simulate OOXML run splitting: "ACME Corp" | " Ltd."
        para.add_run("ACME Corp")
        para.add_run(" Ltd.")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("ACME Corp Ltd.", "ACME Inc Ltd.", author="Tester")

        assert count == 1
        assert any(d.text == "ACME Corp Ltd." for d in rp.deletions)
        assert any(i.text == "ACME Inc Ltd." for i in rp.insertions)

    def it_finds_text_spanning_many_runs(self):
        """Date split across 7 runs like real OOXML: 20/04/2027."""
        doc = Document()
        para = doc.add_paragraph("")
        for chunk in ["20", "/", "0", "4", "/202", "7"]:
            para.add_run(chunk)
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("20/04/2027", "31/12/2027", author="Tester")

        assert count == 1
        assert any(d.text == "20/04/2027" for d in rp.deletions)
        assert any(i.text == "31/12/2027" for i in rp.insertions)

    def it_handles_cross_run_with_surrounding_text(self):
        """Cross-run match with text before and after in the boundary runs."""
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("amount of 26.000")
        para.add_run(" Euros (TWENTY")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("26.000 Euros", "30.000 Euros", author="Tester")

        assert count == 1
        assert any(d.text == "26.000 Euros" for d in rp.deletions)
        assert any(i.text == "30.000 Euros" for i in rp.insertions)
        assert "amount of" in rp.accepted_text
        assert "(TWENTY" in rp.accepted_text

    def it_handles_multiple_cross_run_matches(self):
        """Multiple occurrences where each spans runs."""
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("Hello Wor")
        para.add_run("ld and Wor")
        para.add_run("ld again")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("World", "Earth", author="Tester")

        assert count == 2
        assert len(rp.deletions) == 2
        assert len(rp.insertions) == 2


class DescribeRevisionParagraph_replace_tracked_at:
    """Tests for RevisionParagraph.replace_tracked_at."""

    def it_replaces_within_single_run(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)

        rp.replace_tracked_at(start=6, end=11, replace_text="Universe", author="Tester")

        assert rp.has_track_changes is True
        assert any(d.text == "World" for d in rp.deletions)
        assert any(i.text == "Universe" for i in rp.insertions)

    def it_replaces_across_multiple_runs(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("Hello ")
        para.add_run("World")
        rp = RevisionParagraph.from_paragraph(para)

        # Replace "o Wor" (spans two runs: "Hello " and "World")
        rp.replace_tracked_at(start=4, end=9, replace_text="X", author="Tester")

        assert rp.has_track_changes is True
        assert any(d.text == "o Wor" for d in rp.deletions)
        assert any(i.text == "X" for i in rp.insertions)

    def it_preserves_before_and_after_text(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("Hello ")
        para.add_run("World")
        rp = RevisionParagraph.from_paragraph(para)

        rp.replace_tracked_at(start=4, end=9, replace_text="X", author="Tester")

        accepted = rp.accepted_text
        assert "Hell" in accepted
        assert "X" in accepted
        assert "ld" in accepted

    def it_raises_on_invalid_offsets(self):
        doc = Document()
        para = doc.add_paragraph("Hello")
        rp = RevisionParagraph.from_paragraph(para)

        with pytest.raises(ValueError, match="Invalid offsets"):
            rp.replace_tracked_at(start=10, end=15, replace_text="test", author="Tester")

    def it_raises_on_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("")
        rp = RevisionParagraph.from_paragraph(para)

        with pytest.raises(ValueError, match="Invalid offsets"):
            rp.replace_tracked_at(start=0, end=5, replace_text="test", author="Tester")


class DescribeRevisionDocument_find_and_replace_tracked:
    """Tests for RevisionDocument.find_and_replace_tracked."""

    def it_replaces_across_entire_document(self):
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("World is great")
        doc.add_paragraph("No match")
        rdoc = RevisionDocument(doc)

        count = rdoc.find_and_replace_tracked("World", "Universe", author="Bot")

        assert count == 2
        changes = rdoc.track_changes
        assert len(changes) > 0

    def it_returns_zero_for_no_matches(self):
        doc = Document()
        doc.add_paragraph("Hello")
        rdoc = RevisionDocument(doc)

        count = rdoc.find_and_replace_tracked("zzz", "yyy", author="Bot")

        assert count == 0

    def it_supports_accepted_index_mode(self):
        doc = Document()
        para = doc.add_paragraph("Before ")
        rp = RevisionParagraph.from_paragraph(para)
        # Prior tracked insertion — visible only in accepted_text
        rp.add_tracked_insertion("Unisys", author="Prior")
        rdoc = RevisionDocument(doc)

        count = rdoc.find_and_replace_tracked("Unisys", "NewCo", author="Bot", index_mode="accepted")

        assert count == 1


class DescribeRevisionParagraph_accepted_index_mode:
    """Tests for index_mode='accepted' on replace / delete methods."""

    def _para_with_prior_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Hello ")  # "Hello " in a w:r
        rp = RevisionParagraph.from_paragraph(para)
        rp.add_tracked_insertion("Unisys World", author="Prior")  # inside w:ins
        return rp

    def _para_with_prior_deletion(self):
        doc = Document()
        para = doc.add_paragraph("Hello DELETED World")
        rp = RevisionParagraph.from_paragraph(para)
        # Mark "DELETED " as a tracked deletion (offsets into raw text)
        rp.add_tracked_deletion(start=6, end=14, author="Prior")
        return rp

    def it_finds_text_inside_prior_insertion_with_accepted_mode(self):
        rp = self._para_with_prior_insertion()
        # accepted_text == "Hello Unisys World"; raw self.text == "Hello "
        count = rp.replace_tracked("Unisys", "NewCo", author="Bot", index_mode="accepted")

        assert count == 1
        # New w:del/w:ins are nested inside the prior w:ins — search descendants.
        del_texts = ["".join(t.text or "" for t in d.xpath(".//w:delText")) for d in rp._p.xpath(".//w:del")]
        ins_texts = ["".join(t.text or "" for t in i.xpath("./w:r/w:t")) for i in rp._p.xpath(".//w:ins")]
        assert "Unisys" in del_texts
        assert "NewCo" in ins_texts
        # Round-trip: accepted view should now read "Hello NewCo World"
        assert rp.accepted_text == "Hello NewCo World"

    def it_default_mode_cannot_find_text_inside_prior_insertion(self):
        rp = self._para_with_prior_insertion()
        count = rp.replace_tracked("Unisys", "NewCo", author="Bot")
        assert count == 0

    def it_offsets_skip_deleted_content_in_accepted_mode(self):
        rp = self._para_with_prior_deletion()
        # accepted_text == "Hello World" (the deletion is hidden)
        assert rp.accepted_text == "Hello World"
        # Replace "World" at accepted offsets [6, 11)
        rp.replace_tracked_at(start=6, end=11, replace_text="Earth", author="Bot", index_mode="accepted")

        assert any(d.text == "World" for d in rp.deletions)
        assert any(i.text == "Earth" for i in rp.insertions)

    def it_original_mode_sees_deleted_but_not_inserted(self):
        doc = Document()
        para = doc.add_paragraph("Keep DEL ")
        rp = RevisionParagraph.from_paragraph(para)
        rp.add_tracked_deletion(start=5, end=8, author="Prior")  # "DEL"
        rp.add_tracked_insertion("INS", author="Prior")

        # original_text includes "DEL" (kept) but excludes "INS" (skipped)
        assert "DEL" in rp.original_text
        assert "INS" not in rp.original_text
        # accepted_text excludes "DEL" but includes "INS"
        assert "DEL" not in rp.accepted_text
        assert "INS" in rp.accepted_text

    def it_add_tracked_deletion_with_accepted_mode(self):
        rp = self._para_with_prior_insertion()
        # accepted_text == "Hello Unisys World" — delete "Unisys" (6..12)
        rp.add_tracked_deletion(start=6, end=12, author="Bot", index_mode="accepted")

        # The deletion should now exist somewhere (nested inside the w:ins)
        # and rendering accepted_text should no longer contain "Unisys"
        assert "Unisys" not in rp.accepted_text

    def it_round_trips_through_accept_all(self):
        rp = self._para_with_prior_insertion()
        doc = rp._parent._parent  # underlying python-docx document
        rdoc = RevisionDocument(doc)
        rdoc.find_and_replace_tracked("Unisys", "NewCo", author="Bot", index_mode="accepted")
        rdoc.accept_all()

        assert rdoc.paragraphs[0].text == "Hello NewCo World"

    def it_default_mode_behaves_unchanged_for_plain_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("World", "Earth", author="Bot")

        assert count == 1
        assert any(d.text == "World" for d in rp.deletions)
        assert any(i.text == "Earth" for i in rp.insertions)
