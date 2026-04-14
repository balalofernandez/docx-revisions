"""Tests for accepting and rejecting tracked changes."""

from pathlib import Path

from docx import Document

import docx_revisions  # noqa: F401
from docx_revisions import RevisionDocument, RevisionParagraph


class DescribeAcceptReject_individual:
    """Accept/reject individual tracked changes."""

    def it_accepts_an_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Before ")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_insertion(text="inserted", author="A", revision_id=1)

        tracked.accept()

        assert rp.has_track_changes is False
        assert "inserted" in rp.text

    def it_rejects_an_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Before ")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_insertion(text="inserted", author="A", revision_id=1)

        tracked.reject()

        assert rp.has_track_changes is False
        assert "inserted" not in rp.text

    def it_accepts_a_deletion(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_deletion(start=0, end=5, author="A", revision_id=1)

        tracked.accept()

        assert rp.has_track_changes is False
        # After accepting deletion, "Hello" should be gone
        assert "Hello" not in rp.text

    def it_rejects_a_deletion(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_deletion(start=0, end=5, author="A", revision_id=1)

        tracked.reject()

        assert rp.has_track_changes is False
        # After rejecting deletion, "Hello" should be restored as normal text
        assert "Hello" in rp.text


class DescribeAcceptReject_document_level:
    """Accept/reject all changes in a document."""

    def it_accepts_all_changes(self):
        doc = Document()
        p1 = doc.add_paragraph("Start ")
        rp1 = RevisionParagraph.from_paragraph(p1)
        rp1.add_tracked_insertion(text="added", author="A", revision_id=1)

        p2 = doc.add_paragraph("Keep this")
        rp2 = RevisionParagraph.from_paragraph(p2)
        rp2.add_tracked_deletion(start=0, end=4, author="A", revision_id=2)

        rdoc = RevisionDocument(doc)
        rdoc.accept_all()

        # After accepting all: insertions kept, deletions removed
        assert len(rdoc.track_changes) == 0

    def it_rejects_all_changes(self):
        doc = Document()
        p1 = doc.add_paragraph("Start ")
        rp1 = RevisionParagraph.from_paragraph(p1)
        rp1.add_tracked_insertion(text="added", author="A", revision_id=1)

        rdoc = RevisionDocument(doc)
        rdoc.reject_all()

        assert len(rdoc.track_changes) == 0


class DescribeAcceptReject_from_docx:
    """Accept/reject using a real docx file."""

    def it_can_accept_all_on_test_docx(self, oxml_docx_path: Path):
        rdoc = RevisionDocument(str(oxml_docx_path))

        # Verify there are track changes first
        assert len(rdoc.track_changes) > 0

        rdoc.accept_all()

        assert len(rdoc.track_changes) == 0


def _build_doc_with_table_changes():
    """Build a Document with tracked changes in body, a table cell, and a nested table."""
    doc = Document()

    body_p = doc.add_paragraph("Body ")
    RevisionParagraph.from_paragraph(body_p).add_tracked_insertion(text="body_ins", author="A", revision_id=1)

    table = doc.add_table(rows=1, cols=2)

    cell_p = table.cell(0, 0).paragraphs[0]
    cell_p.add_run("Cell text here")
    RevisionParagraph.from_paragraph(cell_p).add_tracked_deletion(start=0, end=4, author="A", revision_id=2)

    extra_cell_p = table.cell(0, 1).add_paragraph("Second cell ")
    RevisionParagraph.from_paragraph(extra_cell_p).add_tracked_insertion(text="cell_ins", author="A", revision_id=3)

    # Nested table inside the first cell
    nested_cell = table.cell(0, 0)
    nested = nested_cell.add_table(rows=1, cols=1)
    nested_p = nested.cell(0, 0).paragraphs[0]
    nested_p.add_run("Nested")
    RevisionParagraph.from_paragraph(nested_p).add_tracked_insertion(text="nested_ins", author="A", revision_id=4)

    return doc


class DescribeTables:
    """Accept/reject tracked changes inside tables (issue #9)."""

    def it_sees_tracked_changes_inside_tables(self):
        doc = _build_doc_with_table_changes()
        rdoc = RevisionDocument(doc)

        # All 4 tracked changes should be visible (body + 2 cells + nested)
        assert len(rdoc.track_changes) == 4

    def it_all_paragraphs_includes_table_paragraphs(self):
        doc = _build_doc_with_table_changes()
        rdoc = RevisionDocument(doc)

        # paragraphs (body-only) is smaller than all_paragraphs
        assert len(rdoc.paragraphs) < len(rdoc.all_paragraphs)
        # Body has 1 paragraph; tables contribute at least 3 more
        assert len(rdoc.all_paragraphs) >= 4

    def it_paragraphs_property_remains_body_only(self):
        doc = _build_doc_with_table_changes()
        rdoc = RevisionDocument(doc)

        # Backwards compat: paragraphs returns only body paragraphs
        assert len(rdoc.paragraphs) == 1

    def it_accepts_all_changes_inside_tables(self):
        doc = _build_doc_with_table_changes()
        rdoc = RevisionDocument(doc)

        rdoc.accept_all()

        assert len(rdoc.track_changes) == 0
        for para in rdoc.all_paragraphs:
            assert para.has_track_changes is False

    def it_rejects_all_changes_inside_tables(self):
        doc = _build_doc_with_table_changes()
        rdoc = RevisionDocument(doc)

        rdoc.reject_all()

        assert len(rdoc.track_changes) == 0
        for para in rdoc.all_paragraphs:
            assert para.has_track_changes is False

    def it_accept_all_preserves_insertion_text_in_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell_p = table.cell(0, 0).paragraphs[0]
        cell_p.add_run("Start ")
        RevisionParagraph.from_paragraph(cell_p).add_tracked_insertion(text="kept", author="A", revision_id=1)

        rdoc = RevisionDocument(doc)
        rdoc.accept_all()

        final_text = RevisionParagraph.from_paragraph(rdoc.document.tables[0].cell(0, 0).paragraphs[0]).text
        assert "kept" in final_text

    def it_reject_all_removes_insertion_in_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell_p = table.cell(0, 0).paragraphs[0]
        cell_p.add_run("Start ")
        RevisionParagraph.from_paragraph(cell_p).add_tracked_insertion(text="gone", author="A", revision_id=1)

        rdoc = RevisionDocument(doc)
        rdoc.reject_all()

        final_text = RevisionParagraph.from_paragraph(rdoc.document.tables[0].cell(0, 0).paragraphs[0]).text
        assert "gone" not in final_text

    def it_find_and_replace_tracked_works_in_nested_tables(self):
        doc = Document()
        doc.add_paragraph("Replace me here")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).paragraphs[0].add_run("Replace me in cell")
        nested = table.cell(0, 0).add_table(rows=1, cols=1)
        nested.cell(0, 0).paragraphs[0].add_run("Replace me nested")

        rdoc = RevisionDocument(doc)
        count = rdoc.find_and_replace_tracked("Replace me", "Done", author="A")

        assert count == 3
