"""RevisionDocument — the main entry point for working with tracked changes.

Opens a ``.docx`` file (or wraps an existing ``Document``) and exposes
paragraphs as ``RevisionParagraph`` objects, plus document-level operations
for accepting, rejecting, and performing find-and-replace with tracking.
"""

from __future__ import annotations

from pathlib import Path
from typing import IO, Iterator, List

from docx import Document as _new_document
from docx.document import Document as _DocumentClass
from docx.table import Table as _Table

from docx_revisions.paragraph import IndexMode, RevisionParagraph
from docx_revisions.revision import TrackedChange


class RevisionDocument:
    """Entry point for reading and writing tracked changes in a docx file.

    Example:
        ```python
        from docx_revisions import RevisionDocument

        rdoc = RevisionDocument("contract.docx")
        for para in rdoc.paragraphs:
            if para.has_track_changes:
                print(para.accepted_text)

        rdoc.accept_all()
        rdoc.save("contract_clean.docx")
        ```
    """

    def __init__(self, path_or_doc: str | Path | IO[bytes] | _DocumentClass | None = None):
        """Open a docx file or wrap an existing ``Document``.

        Args:
            path_or_doc: A file path, file-like object, or an existing
                ``Document`` instance.  Pass ``None`` to create a new
                blank document.
        """
        if isinstance(path_or_doc, _DocumentClass):
            self._document = path_or_doc
        else:
            self._document = _new_document(path_or_doc)

    @property
    def document(self) -> _DocumentClass:
        """The underlying python-docx ``Document`` object."""
        return self._document

    @property
    def paragraphs(self) -> List[RevisionParagraph]:
        """All body paragraphs as ``RevisionParagraph`` objects.

        Only paragraphs in the document body are returned. Paragraphs inside
        tables are excluded. Use :attr:`all_paragraphs` to iterate over every
        paragraph including those nested in tables.
        """
        return [RevisionParagraph.from_paragraph(p) for p in self._document.paragraphs]

    @property
    def all_paragraphs(self) -> List[RevisionParagraph]:
        """Every paragraph in the document, including those inside tables.

        Walks the document body and recurses into all tables (including
        nested tables within cells).

        Example:
            ```python
            rdoc = RevisionDocument("contract.docx")
            for para in rdoc.all_paragraphs:
                if para.has_track_changes:
                    print(para.accepted_text)
            ```
        """
        return list(self._iter_all_paragraphs())

    def _iter_all_paragraphs(self) -> Iterator[RevisionParagraph]:
        """Yield every ``RevisionParagraph`` in the body and all tables.

        Recurses into nested tables via ``cell.tables``.
        """
        for p in self._document.paragraphs:
            yield RevisionParagraph.from_paragraph(p)
        for table in self._document.tables:
            yield from self._iter_table_paragraphs(table)

    def _iter_table_paragraphs(self, table: _Table) -> Iterator[RevisionParagraph]:
        """Yield every ``RevisionParagraph`` inside *table* recursively."""
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield RevisionParagraph.from_paragraph(p)
                for nested in cell.tables:
                    yield from self._iter_table_paragraphs(nested)

    @property
    def track_changes(self) -> List[TrackedChange]:
        """All tracked changes across the document body and tables."""
        changes: List[TrackedChange] = []
        for para in self._iter_all_paragraphs():
            changes.extend(para.track_changes)
        return changes

    def accept_all(self) -> None:
        """Accept every tracked change in the document.

        Insertions are kept (wrapper removed), deletions are removed entirely.
        Tracked changes inside tables (including nested tables) are processed.
        Loops until no tracked changes remain so that nested revisions (which
        can arise from ``replace_tracked(index_mode="accepted")``) are fully
        resolved.
        """
        for para in self._iter_all_paragraphs():
            while para.track_changes:
                for change in list(para.track_changes):
                    change.accept()

    def reject_all(self) -> None:
        """Reject every tracked change in the document.

        Insertions are removed entirely, deletions are kept (wrapper removed,
        ``w:delText`` converted back to ``w:t``). Tracked changes inside
        tables (including nested tables) are processed. Loops until no tracked
        changes remain.
        """
        for para in self._iter_all_paragraphs():
            while para.track_changes:
                for change in list(para.track_changes):
                    change.reject()

    def find_and_replace_tracked(
        self,
        search_text: str,
        replace_text: str,
        author: str = "",
        comment: str | None = None,
        index_mode: IndexMode = "text",
    ) -> int:
        """Find and replace across the whole document with track changes.

        Searches all paragraphs in the document body and tables (including
        nested tables).

        Args:
            search_text: Text to find.
            replace_text: Replacement text.
            author: Author name for the revisions.
            comment: Optional comment text (requires python-docx comment
                support).
            index_mode: Which text view to search against per paragraph.  See
                :meth:`RevisionParagraph.replace_tracked` — ``"text"`` (default),
                ``"accepted"``, or ``"original"``.

        Returns:
            Total number of replacements made.

        Example:
            ```python
            rdoc = RevisionDocument("doc.docx")
            # Replace against the accepted view so matches inside prior
            # tracked insertions are also found.
            count = rdoc.find_and_replace_tracked(
                "Acme Corp", "NewCo Inc", author="Legal", index_mode="accepted"
            )
            rdoc.save("doc_revised.docx")
            ```
        """
        total_count = 0
        for para in self._iter_all_paragraphs():
            total_count += para.replace_tracked(
                search_text, replace_text, author=author, comment=comment, index_mode=index_mode
            )
        return total_count

    def save(self, path_or_stream: str | Path | IO[bytes]) -> None:
        """Save the document to a path or file-like object.

        Args:
            path_or_stream: Destination file path (``str`` or ``Path``) or a
                writable binary file-like object (anything with a ``write``
                method, such as ``io.BytesIO``).

        Raises:
            TypeError: If *path_or_stream* is neither a path nor a writable
                binary stream.
            ValueError: If *path_or_stream* is an empty string, or is a text-
                mode file object.

        Example:
            ```python
            import io
            from docx_revisions import RevisionDocument

            rdoc = RevisionDocument("contract.docx")
            rdoc.accept_all()

            buffer = io.BytesIO()
            rdoc.save(buffer)
            buffer.seek(0)
            data = buffer.read()
            ```
        """
        if isinstance(path_or_stream, str | Path):
            path_str = str(path_or_stream)
            if not path_str:
                raise ValueError("save() path must not be empty")
            self._document.save(path_str)
            return

        write = getattr(path_or_stream, "write", None)
        if not callable(write):
            raise TypeError(
                f"save() expects a str, Path, or writable binary file-like object; got {type(path_or_stream).__name__}"
            )

        mode = getattr(path_or_stream, "mode", None)
        if isinstance(mode, str) and "b" not in mode:
            raise ValueError(
                f"save() requires a binary-mode stream; got mode={mode!r}. "
                "Open the file with mode='wb' or use io.BytesIO()."
            )

        self._document.save(path_or_stream)
