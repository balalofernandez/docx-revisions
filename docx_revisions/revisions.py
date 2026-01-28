"""Core track changes functionality for reading and writing revisions."""

import contextlib
from datetime import datetime
from typing import Iterator

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

from docx_revisions.models import Revision

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _iter_paragraph_content(paragraph: Paragraph) -> Iterator[tuple[str, etree._Element]]:
    """Iterate over paragraph content elements yielding (tag_type, element) tuples.

    Args:
        paragraph: The paragraph to iterate over.
    Yields:
        Tuples of (tag_type, element) where tag_type is 'run', 'ins', or 'del'.
    """
    for element in paragraph._element.iterchildren():
        tag = etree.QName(element.tag).localname
        if tag == "r":
            yield ("run", element)
        elif tag == "ins":
            yield ("ins", element)
        elif tag == "del":
            yield ("del", element)


def _extract_text_from_element(element: etree._Element) -> str:
    """Extract all text content from an element, including nested runs.

    Args:
        element: The XML element to extract text from.
    Returns:
        Concatenated text content.
    """
    texts = []
    for t_elem in element.iter(qn("w:t")):
        if t_elem.text:
            texts.append(t_elem.text)
    # Also check for delText elements (used in deletions)
    for dt_elem in element.iter(qn("w:delText")):
        if dt_elem.text:
            texts.append(dt_elem.text)
    return "".join(texts)


def get_accepted_text(paragraph: Paragraph) -> str:
    """Get text from paragraph with track changes applied.

    Insertions are included, deletions are excluded - as if all changes were accepted.

    Args:
        paragraph: The paragraph to extract text from.
    Returns:
        The accepted text with insertions applied and deletions removed.

    Example:
        >>> from docx import Document
        >>> doc = Document("tracked.docx")
        >>> text = get_accepted_text(doc.paragraphs[0])
    """
    result = []
    for tag_type, element in _iter_paragraph_content(paragraph):
        if tag_type in ("run", "ins"):
            result.append(_extract_text_from_element(element))
        # Skip 'del' elements - they're not part of accepted text
    return "".join(result)


def get_revisions(paragraph: Paragraph) -> list[Revision]:
    """Extract all revisions (insertions and deletions) from a paragraph.

    Args:
        paragraph: The paragraph to extract revisions from.
    Returns:
        List of Revision objects with type, text, author, and date.

    Example:
        >>> from docx import Document
        >>> doc = Document("tracked.docx")
        >>> revisions = get_revisions(doc.paragraphs[0])
        >>> for rev in revisions:
        ...     print(f"{rev.type}: '{rev.text}' by {rev.author}")
    """
    revisions = []
    for tag_type, element in _iter_paragraph_content(paragraph):
        if tag_type in ("ins", "del"):
            author = element.get(qn("w:author"), "Unknown")
            date_str = element.get(qn("w:date"))

            date = None
            if date_str:
                with contextlib.suppress(ValueError, AttributeError):
                    date = datetime.fromisoformat(date_str.replace("Z", "+00:00"))

            text = _extract_text_from_element(element)
            if text:
                revisions.append(Revision(type=tag_type, text=text, author=author, date=date))

    return revisions


def _create_run_with_text(text: str) -> etree._Element:
    """Create a w:r element containing the given text.

    Args:
        text: The text content for the run.
    Returns:
        A new w:r element with w:t child containing the text.
    """
    run = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    # Preserve spaces
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t_elem)
    return run


def _create_del_run_with_text(text: str) -> etree._Element:
    """Create a w:r element with w:delText for deleted content.

    Args:
        text: The deleted text content.
    Returns:
        A new w:r element with w:delText child.
    """
    run = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.text = text
    del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(del_text)
    return run


def insert_with_tracking(
    paragraph: Paragraph, text: str, position: int = -1, author: str = "Agent", date: datetime | None = None
) -> None:
    """Insert text with track changes at the specified position.

    Args:
        paragraph: The paragraph to insert into.
        text: The text to insert.
        position: Character position to insert at. -1 means append at end.
        author: Author name for the revision metadata.
        date: Timestamp for the revision. Defaults to current time.

    Example:
        >>> from docx import Document
        >>> doc = Document()
        >>> p = doc.add_paragraph("Hello world")
        >>> insert_with_tracking(p, " beautiful", position=5, author="Editor")
    """
    if date is None:
        date = datetime.now()

    ins_elem = OxmlElement("w:ins")
    ins_elem.set(qn("w:author"), author)
    ins_elem.set(qn("w:date"), date.isoformat())
    ins_elem.set(qn("w:id"), "0")

    run = _create_run_with_text(text)
    ins_elem.append(run)

    if position == -1:
        paragraph._element.append(ins_elem)
    else:
        _insert_at_position(paragraph, ins_elem, position)


def _insert_at_position(paragraph: Paragraph, new_element: etree._Element, position: int) -> None:
    """Insert an element at a specific character position in the paragraph.

    Args:
        paragraph: The paragraph to modify.
        new_element: The element to insert.
        position: The character offset where insertion should occur.
    """
    current_pos = 0
    for child in list(paragraph._element):
        tag = etree.QName(child.tag).localname
        if tag in ("r", "ins"):
            text = _extract_text_from_element(child)
            text_len = len(text)

            if current_pos + text_len > position:
                # Split this element
                offset_in_element = position - current_pos
                _split_and_insert(paragraph._element, child, new_element, offset_in_element)
                return

            current_pos += text_len
        elif tag == "del":
            continue  # Don't count deleted text

    # Position is at or past end, just append
    paragraph._element.append(new_element)


def _split_and_insert(parent: etree._Element, target: etree._Element, new_element: etree._Element, offset: int) -> None:
    """Split a run/ins element at offset and insert new_element between the halves.

    Args:
        parent: The parent element (paragraph).
        target: The element to split.
        new_element: The element to insert at the split point.
        offset: Character offset within target where split occurs.
    """
    text = _extract_text_from_element(target)
    before_text = text[:offset]
    after_text = text[offset:]

    target_index = list(parent).index(target)
    tag = etree.QName(target.tag).localname

    # Remove original
    parent.remove(target)

    # Create before element
    if before_text:
        if tag == "ins":
            before_elem = OxmlElement("w:ins")
            for attr in target.attrib:
                before_elem.set(attr, target.get(attr))
            before_elem.append(_create_run_with_text(before_text))
        else:
            before_elem = _create_run_with_text(before_text)
        parent.insert(target_index, before_elem)
        target_index += 1

    # Insert new element
    parent.insert(target_index, new_element)
    target_index += 1

    # Create after element
    if after_text:
        if tag == "ins":
            after_elem = OxmlElement("w:ins")
            for attr in target.attrib:
                after_elem.set(attr, target.get(attr))
            after_elem.append(_create_run_with_text(after_text))
        else:
            after_elem = _create_run_with_text(after_text)
        parent.insert(target_index, after_elem)


def delete_with_tracking(
    paragraph: Paragraph, start_offset: int, end_offset: int, author: str = "Agent", date: datetime | None = None
) -> None:
    """Mark a text range as deleted with track changes.

    Args:
        paragraph: The paragraph to delete from.
        start_offset: Character offset where deletion starts (inclusive).
        end_offset: Character offset where deletion ends (exclusive).
        author: Author name for the revision metadata.
        date: Timestamp for the revision. Defaults to current time.

    Example:
        >>> from docx import Document
        >>> doc = Document()
        >>> p = doc.add_paragraph("Hello beautiful world")
        >>> delete_with_tracking(p, 6, 16, author="Editor")  # Deletes "beautiful "
    """
    if date is None:
        date = datetime.now()

    if start_offset >= end_offset:
        return

    # Find and mark text in range as deleted
    _mark_range_as_deleted(paragraph, start_offset, end_offset, author, date)


def _mark_range_as_deleted(paragraph: Paragraph, start: int, end: int, author: str, date: datetime) -> None:
    """Mark text range as deleted by wrapping it in w:del elements.

    Args:
        paragraph: The paragraph to modify.
        start: Start character offset.
        end: End character offset.
        author: Author for deletion metadata.
        date: Timestamp for deletion.
    """
    current_pos = 0
    children = list(paragraph._element)
    new_children = []

    for child in children:
        tag = etree.QName(child.tag).localname

        if tag == "del":
            new_children.append(child)
            continue

        if tag not in ("r", "ins"):
            new_children.append(child)
            continue

        text = _extract_text_from_element(child)
        text_len = len(text)
        elem_start = current_pos
        elem_end = current_pos + text_len

        # Check overlap with deletion range
        if elem_end <= start or elem_start >= end:
            # No overlap
            new_children.append(child)
        elif elem_start >= start and elem_end <= end:
            # Fully contained - wrap entire element in del
            del_elem = _create_deletion_element(text, author, date)
            new_children.append(del_elem)
        else:
            # Partial overlap - need to split
            parts = _split_for_deletion(text, elem_start, start, end, author, date, tag, child)
            new_children.extend(parts)

        current_pos += text_len

    # Replace children
    for child in list(paragraph._element):
        paragraph._element.remove(child)
    for child in new_children:
        paragraph._element.append(child)


def _create_deletion_element(text: str, author: str, date: datetime) -> etree._Element:
    """Create a w:del element wrapping the deleted text.

    Args:
        text: The deleted text.
        author: Author of the deletion.
        date: Timestamp of the deletion.
    Returns:
        A w:del element containing the text.
    """
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), date.isoformat())
    del_elem.set(qn("w:id"), "0")
    del_elem.append(_create_del_run_with_text(text))
    return del_elem


def _split_for_deletion(
    text: str,
    elem_start: int,
    del_start: int,
    del_end: int,
    author: str,
    date: datetime,
    tag: str,
    original: etree._Element,
) -> list[etree._Element]:
    """Split text for partial deletion, returning list of elements.

    Args:
        text: The full text of the element.
        elem_start: Character position where this element starts.
        del_start: Start of deletion range.
        del_end: End of deletion range.
        author: Author for deletion metadata.
        date: Timestamp for deletion.
        tag: Original element tag ('r' or 'ins').
        original: Original element (for copying attributes if ins).
    Returns:
        List of elements representing the split.
    """
    result = []
    local_del_start = max(0, del_start - elem_start)
    local_del_end = min(len(text), del_end - elem_start)

    before = text[:local_del_start]
    deleted = text[local_del_start:local_del_end]
    after = text[local_del_end:]

    if before:
        if tag == "ins":
            elem = OxmlElement("w:ins")
            for attr in original.attrib:
                elem.set(attr, original.get(attr))
            elem.append(_create_run_with_text(before))
        else:
            elem = _create_run_with_text(before)
        result.append(elem)

    if deleted:
        result.append(_create_deletion_element(deleted, author, date))

    if after:
        if tag == "ins":
            elem = OxmlElement("w:ins")
            for attr in original.attrib:
                elem.set(attr, original.get(attr))
            elem.append(_create_run_with_text(after))
        else:
            elem = _create_run_with_text(after)
        result.append(elem)

    return result


def replace_with_tracking(
    paragraph: Paragraph, old_text: str, new_text: str, author: str = "Agent", date: datetime | None = None
) -> bool:
    """Replace text with track changes (deletion + insertion).

    Finds the first occurrence of old_text and replaces it with new_text,
    marking the old text as deleted and the new text as inserted.

    Args:
        paragraph: The paragraph to modify.
        old_text: The text to find and replace.
        new_text: The replacement text.
        author: Author name for the revision metadata.
        date: Timestamp for the revision. Defaults to current time.
    Returns:
        True if replacement was made, False if old_text was not found.

    Example:
        >>> from docx import Document
        >>> doc = Document()
        >>> p = doc.add_paragraph("Hello world")
        >>> replace_with_tracking(p, "world", "universe", author="Editor")
        True
    """
    if date is None:
        date = datetime.now()

    # Find position of old_text in accepted text
    accepted = get_accepted_text(paragraph)
    pos = accepted.find(old_text)

    if pos == -1:
        return False

    # Delete old text and insert new text at same position
    delete_with_tracking(paragraph, pos, pos + len(old_text), author, date)
    insert_with_tracking(paragraph, new_text, pos, author, date)

    return True
